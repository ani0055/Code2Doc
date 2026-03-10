from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import re
import base64
from datetime import datetime
from PIL import Image as PILImage



class ExportService:
    """
    Export service using WeasyPrint for PDF and python-docx for DOCX.
    Supports embedding diagram images.
    """
    
    def generate_pdf(self, markdown_text: str, filename: str = "documentation", diagram_image: str = None) -> bytes:
        """
        Generate PDF using WeasyPrint with optional diagram image.
        """
        try:
            html_content = self._create_html_document(markdown_text, filename, diagram_image)
            pdf_bytes = HTML(string=html_content).write_pdf()
            return pdf_bytes
            
        except Exception as e:
            raise Exception(f"PDF generation failed: {str(e)}")
    
    def _create_html_document(self, markdown_text: str, filename: str = "documentation", diagram_image: str = None) -> str:
        """
        Create a complete HTML document with embedded CSS and optional diagram image.
        """
        html_body = self._markdown_to_html(markdown_text)
        
        # Add diagram section if provided
        diagram_section = ""
        if diagram_image:
            try:
                # Process base64 image
                if diagram_image.startswith('data:image'):
                    diagram_image = diagram_image.split(',')[1]
                
                # Embed image directly in HTML
                diagram_section = f"""
                <div class="page-break"></div>
                <h2>Code Flow Diagram</h2>
                <div class="diagram-container">
                    <img src="data:image/png;base64,{diagram_image}" alt="Code Flow Diagram" class="diagram-image"/>
                </div>
                """
            except Exception as e:
                print(f"Failed to process diagram image: {e}")
                diagram_section = ""
        
        html_document = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{filename}</title>
    <style>
        @page {{
            size: A4;
            margin: 2cm;
            @bottom-right {{
                content: "Page " counter(page) " of " counter(pages);
                font-size: 9pt;
                color: #666;
            }}
            @bottom-left {{
                content: "Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}";
                font-size: 9pt;
                color: #666;
            }}
        }}
        
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Times New Roman', Times, serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #1f2937;
            margin: 0;
            padding: 0;
            max-width: 100%;
            word-wrap: break-word;
        }}
        
        h1 {{
            font-size: 24pt;
            color: #1f2937;
            text-align: center;
            border-bottom: 3px solid #3498db;
            padding: 0;
            margin: 0;
            page-break-before: avoid;
            page-break-after: avoid;
        }}
        
        h2 {{
            font-size: 16pt;
            color: #374151;
            margin-top: 20pt;
            margin-bottom: 12pt;
            border-left: 4px solid #3498db;
            padding-left: 12pt;
            page-break-after: avoid;
        }}
        
        h3 {{
            font-size: 13pt;
            color: #4b5563;
            margin-top: 16pt;
            margin-bottom: 8pt;
            page-break-after: avoid;
        }}
        
        h4 {{
            font-size: 12pt;
            color: #6b7280;
            margin-top: 12pt;
            margin-bottom: 6pt;
            page-break-after: avoid;
        }}
        
        p {{
            margin-bottom: 10pt;
            text-align: justify;
            orphans: 3;
            widows: 3;
        }}
        
        ul, ol {{
            margin-left: 25pt;
            margin-bottom: 10pt;
        }}
        
        li {{
            margin-bottom: 6pt;
        }}
        
        strong {{
            font-weight: bold;
            color: #1f2937;
        }}
        
        code {{
            font-family: 'Courier New', Courier, monospace;
            font-size: 9pt;
            background-color: #f3f4f6;
            color: #059669;
            padding: 2pt 4pt;
            border-radius: 3pt;
        }}
        
        pre {{
            background-color: #f8f9fa;
            border: 1pt solid #e5e7eb;
            border-left: 4pt solid #3498db;
            padding: 12pt;
            margin: 12pt 0;
            overflow-x: auto;
            page-break-inside: avoid;
            border-radius: 4pt;
        }}
        
        pre code {{
            background: none;
            padding: 0;
            font-size: 9pt;
            color: #059669;
            display: block;
            white-space: pre-wrap;
            word-wrap: break-word;
        }}
        
        .page-break {{
            page-break-before: always;
        }}
        
        .diagram-container {{
            margin: 20pt 0;
            text-align: center;
            page-break-inside: avoid;
        }}
        
        .diagram-image {{
            max-width: 100%;
            height: auto;
            border: 1pt solid #e5e7eb;
            border-radius: 4pt;
            box-shadow: 0 2pt 4pt rgba(0,0,0,0.1);
        }}
        
        .info-box {{
            background-color: #e8f4f8;
            border: 1pt solid #b8dae6;
            padding: 10pt;
            margin: 10pt 0;
            border-radius: 4pt;
        }}
        
        .info-box p {{
            margin: 0;
        }}
        
        a {{
            color: #3498db;
            text-decoration: none;
        }}
        
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15pt 0;
            page-break-inside: avoid;
        }}
        
        th, td {{
            border: 1pt solid #e5e7eb;
            padding: 8pt;
            text-align: left;
        }}
        
        th {{
            background-color: #3498db;
            color: white;
            font-weight: bold;
        }}
        
        tr:nth-child(even) {{
            background-color: #f8f9fa;
        }}
    </style>
</head>
<body>
    <h1>{filename}</h1>
    
    {diagram_section}
    
    <h2>Documentation</h2>
    {html_body}
</body>
</html>
"""
        return html_document
    
    def _markdown_to_html(self, markdown_text: str) -> str:
        """Convert markdown to HTML with proper formatting."""
        html_lines = []
        lines = markdown_text.split('\n')
        
        in_code_block = False
        code_buffer = []
        in_list = False
        
        for line in lines:
            line_stripped = line.strip()
            
            if line_stripped.startswith('```'):
                if in_code_block:
                    code_content = '\n'.join(code_buffer)
                    html_lines.append(f'<pre><code>{self._escape_html(code_content)}</code></pre>')
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                    if in_list:
                        html_lines.append('</ul>')
                        in_list = False
                continue
            
            if in_code_block:
                code_buffer.append(line)
                continue
            
            if not line_stripped:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                if html_lines and not html_lines[-1].startswith('<p>'):
                    html_lines.append('<p></p>')
                continue
            
            if line_stripped.startswith('#### '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                content = self._apply_inline_formatting_html(line_stripped[5:])
                html_lines.append(f'<h4>{content}</h4>')
            elif line_stripped.startswith('### '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                content = self._apply_inline_formatting_html(line_stripped[4:])
                html_lines.append(f'<h3>{content}</h3>')
            elif line_stripped.startswith('## '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                content = self._apply_inline_formatting_html(line_stripped[3:])
                html_lines.append(f'<h2>{content}</h2>')
            elif line_stripped.startswith('# '):
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                content = self._apply_inline_formatting_html(line_stripped[2:])
                html_lines.append(f'<h1>{content}</h1>')
            elif line.lstrip().startswith(('* ', '- ', '• ')):
                if not in_list:
                    html_lines.append('<ul>')
                    in_list = True
                
                content_with_marker = line.lstrip()
                marker_end_index = content_with_marker.find(' ') + 1
                content = content_with_marker[marker_end_index:]
                content = self._apply_inline_formatting_html(content)
                
                leading_spaces = len(line) - len(line.lstrip())
                padding = f'style="padding-left: {leading_spaces * 0.25}em;"' if leading_spaces > 0 else ''
                
                html_lines.append(f'<li {padding}>{content}</li>')
            elif re.match(r'^\d+\.\s', line_stripped):
                if not in_list:
                    html_lines.append('<ol>')
                    in_list = True
                content = re.sub(r'^\d+\.\s', '', line_stripped)
                content = self._apply_inline_formatting_html(content)
                html_lines.append(f'<li>{content}</li>')
            else:
                if in_list:
                    html_lines.append('</ul>')
                    in_list = False
                content = self._apply_inline_formatting_html(line_stripped)
                html_lines.append(f'<p>{content}</p>')
        
        if in_list:
            html_lines.append('</ul>')
        
        return '\n'.join(html_lines)
    
    def _escape_html(self, text: str) -> str:
        """Escape HTML special characters."""
        if not text:
            return ""
        return (text.replace('&', '&amp;')
                   .replace('<', '&lt;')
                   .replace('>', '&gt;')
                   .replace('"', '&quot;')
                   .replace("'", '&#39;'))
    
    def _apply_inline_formatting_html(self, text: str) -> str:
        """Apply inline formatting (bold and inline code) to HTML."""
        if not text:
            return ""
            
        text = re.sub(r'<\/?strong>', r'**', text, flags=re.IGNORECASE)
        text = re.sub(r'<\/?code>', r'`', text, flags=re.IGNORECASE)
        text = self._escape_html(text)
        text = re.sub(r'\*\*\*(.+?)\*\*\*', r'<strong><em>\1</em></strong>', text)
        text = re.sub(r'(?<!\*)\*\*([^\*]+?)\*\*(?!\*)', r'<strong>\1</strong>', text)
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        
        return text
    
    def generate_docx(self, markdown_text: str, filename: str = "documentation", diagram_image: str = None) -> bytes:
        """
        Generate DOCX using python-docx with optional diagram image.
        """
        try:
            document = Document()
            
            document.core_properties.title = filename
            document.core_properties.author = "Code2Doc"
            document.core_properties.created = datetime.now()
            
            title = document.add_heading(filename, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add diagram if provided
            if diagram_image:
                try:
                    if diagram_image.startswith('data:image'):
                        diagram_image = diagram_image.split(',')[1]
                    
                    image_data = base64.b64decode(diagram_image)
                    image_buffer = io.BytesIO(image_data)
                    
                    # Verify it's a valid image
                    img = PILImage.open(image_buffer)
                    image_buffer.seek(0)
                    
                    document.add_heading('Code Flow Diagram', level=2)
                    
                    # Add image with proper sizing
                    document.add_picture(image_buffer, width=Inches(6))
                    last_paragraph = document.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    document.add_paragraph()  # Spacing
                    
                except Exception as e:
                    print(f"Failed to add diagram to DOCX: {e}")
            
            # Add documentation section
            document.add_heading('Documentation', level=2)
            self._add_markdown_to_docx(document, markdown_text)
            
            buffer = io.BytesIO()
            document.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"DOCX generation failed: {str(e)}")
    
    def _add_markdown_to_docx(self, document, markdown_text: str):
        """Parse markdown and add formatted content to DOCX document."""
        
        markdown_text = re.sub(r'<\/?strong>', r'**', markdown_text, flags=re.IGNORECASE)
        markdown_text = re.sub(r'<\/?code>', r'`', markdown_text, flags=re.IGNORECASE)
        
        lines = markdown_text.split('\n')
        in_code_block = False
        code_buffer = []
        
        for line in lines:
            line_stripped = line.strip()
            
            if line_stripped.startswith('```'):
                if in_code_block:
                    code_text = '\n'.join(code_buffer)
                    code_para = document.add_paragraph(code_text)
                    code_para.style = 'No Spacing'
                    
                    for run in code_para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(5, 150, 105)
                    
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), 'f3f4f6')
                    code_para._element.get_or_add_pPr().append(shading)
                    
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_buffer.append(line)
                continue
            
            if not line_stripped:
                document.add_paragraph()
                continue
            
            if line_stripped.startswith('#### '):
                document.add_heading(line_stripped[5:], level=4)
            elif line_stripped.startswith('### '):
                document.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith('## '):
                document.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith('# '):
                document.add_heading(line_stripped[2:], level=1)
            elif line.lstrip().startswith(('* ', '- ', '• ')):
                leading_spaces = len(line) - len(line.lstrip())
                level = leading_spaces // 4
                
                style = 'List Bullet'
                if level == 1:
                    style = 'List Bullet 2'
                
                para = document.add_paragraph(style=style)
                
                content_with_marker = line.lstrip()
                marker_end_index = content_with_marker.find(' ') + 1
                content = content_with_marker[marker_end_index:]
                
                self._add_formatted_text(para, content)
            elif re.match(r'^\d+\.\s', line_stripped):
                para = document.add_paragraph(style='List Number')
                content = re.sub(r'^\d+\.\s', '', line_stripped)
                self._add_formatted_text(para, content)
            else:
                para = document.add_paragraph()
                self._add_formatted_text(para, line_stripped)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting to a paragraph."""
        if not text:
            return
        
        parts = re.split(r'(\*\*\*[^*]+?\*\*\*|\*\*[^\*]+?\*\*|`[^`]+`)', text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('***') and part.endswith('***') and len(part) > 6:
                run = paragraph.add_run(part[3:-3])
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('`') and part.endswith('`') and len(part) > 2:
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(5, 150, 105)
            else:
                paragraph.add_run(part)